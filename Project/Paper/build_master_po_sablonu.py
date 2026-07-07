# -*- coding: utf-8 -*-
"""Generisanje dokumenta Master_rad_po_sablonu.docx.

Sadržaj rada je u modulima sadrzaj_1.py, sadrzaj_2.py i sadrzaj_3.py, kao
liste blokova. Ova skripta ih renderuje u Word dokument po FTN šablonu:
Times New Roman 12, obostrano poravnanje, ćirilica, naslovi na novim
stranama, slike sa naslovom ispod, tabele sa naslovom iznad, literatura
numerisana po redosledu pojavljivanja u tekstu.

Pokretanje:  python build_master_po_sablonu.py   (iz Project/Paper)

Vrste blokova:
  ('h1', 'Наслов')                      — glavno poglavlje, nova strana
  ('h2', 'Поднаслов') / ('h3', ...)
  ('p', 'текст')                        — pasus; *kurziv*, **bold**, {ref:kljuc}
  ('ul', ['stavka', ...])               — lista sa crticama
  ('img', 'figures/ime.png', sirina_cm, 'Слика X.Y. Наслов')
  ('table', 'Табела X.Y. Наслов', [kolone], [[red], ...], [sirine_cm] | None)
  ('listing', 'Листинг X.Y. Наслов', 'kod')
  ('pagebreak',)
"""

import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from sadrzaj_1 import BLOCKS_1, KDI_SR, KDI_EN
from sadrzaj_2 import BLOCKS_2
from sadrzaj_3 import BLOCKS_3, REFS, BLOCKS_BIOGRAFIJA, BLOCKS_DODACI

OUT_PATH = os.path.join(HERE, 'Master_rad_po_sablonu.docx')

FONT = 'Times New Roman'
MONO = 'Consolas'

# ---------------------------------------------------------------- pomoćno

ref_order = []          # ključevi referenci po redosledu pojavljivanja


def ref_num(key):
    if key not in REFS:
        raise KeyError(f'Nepoznata referenca: {key}')
    if key not in ref_order:
        ref_order.append(key)
    return ref_order.index(key) + 1


_REF_RE = re.compile(r'\{ref:([a-z0-9_]+)\}')


def resolve_refs(text):
    """{ref:a}{ref:b} -> [1][2] ; grupisano {ref:a,b} nije podržano."""
    return _REF_RE.sub(lambda m: f'[{ref_num(m.group(1))}]', text)


_INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)', re.S)


def add_runs(par, text, size=12, base_italic=False):
    """Dodaje tekst u pasus; *kurziv*, **bold**."""
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        bold = chunk.startswith('**') and chunk.endswith('**')
        ital = (not bold) and chunk.startswith('*') and chunk.endswith('*')
        clean = chunk[2:-2] if bold else (chunk[1:-1] if ital else chunk)
        run = par.add_run(clean)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = ital or base_italic
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), FONT)


def para(doc, text='', size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=6, space_before=0, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if text:
        add_runs(p, text, size=size, base_italic=italic)
    return p


def set_style_font(style, name=FONT, size=12, bold=None, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(attr), name)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = FONT
    run.font.size = Pt(11)
    fld1, instr, fld2 = OxmlElement('w:fldChar'), OxmlElement('w:instrText'), OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld2.set(qn('w:fldCharType'), 'end')
    run._element.append(fld1)
    run._element.append(instr)
    run._element.append(fld2)


def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    fld1.set(qn('w:dirty'), 'true')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = 'Садржај се освежава у програму Word: десни клик → Update Field.'
    r2 = OxmlElement('w:r'); r2.append(t)
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._element.append(fld1)
    run._element.append(instr)
    run._element.append(fld2)
    p._element.append(r2)
    p2 = doc.add_paragraph()
    r3 = p2.add_run()
    r3._element.append(fld3)


def set_cell(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             mono=False, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for i, line in enumerate(str(text).split('\n')):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(2)
        if mono:
            run = p.add_run(line)
            run.font.name = MONO
            run.font.size = Pt(size)
            run.italic = italic
        else:
            add_runs(p, line, size=size)
            for r in p.runs:
                r.bold = r.bold or bold
                r.italic = r.italic or italic


def shade_cell(cell, hexcolor):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


# ---------------------------------------------------------------- renderer

def render_blocks(doc, blocks):
    for block in blocks:
        kind = block[0]
        if kind == 'h1':
            doc.add_page_break()
            h = doc.add_heading(resolve_refs(block[1]), level=1)
            for r in h.runs:
                r.font.name = FONT
        elif kind == 'h2':
            h = doc.add_heading(resolve_refs(block[1]), level=2)
            for r in h.runs:
                r.font.name = FONT
        elif kind == 'h3':
            h = doc.add_heading(resolve_refs(block[1]), level=3)
            for r in h.runs:
                r.font.name = FONT
        elif kind == 'p':
            para(doc, resolve_refs(block[1]))
        elif kind == 'ul':
            for item in block[1]:
                p = para(doc, resolve_refs(item), space_after=3)
                p.style = doc.styles['List Bullet']
                p.paragraph_format.left_indent = Cm(0.9)
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(12)
        elif kind == 'img':
            _, path, width, caption = block
            full = os.path.join(HERE, path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.add_run().add_picture(full, width=Cm(width))
            cap = para(doc, resolve_refs(caption), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
            for r in cap.runs:
                r.font.size = Pt(10)
        elif kind == 'table':
            _, caption, header, rows, widths = block
            cap = para(doc, resolve_refs(caption), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
                       space_before=10)
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, txt in enumerate(header):
                set_cell(table.rows[0].cells[j], txt, size=10.5, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
                shade_cell(table.rows[0].cells[j], 'E8EDF4')
            for i, row in enumerate(rows):
                for j, txt in enumerate(row):
                    set_cell(table.rows[i + 1].cells[j], resolve_refs(str(txt)),
                             size=10.5,
                             align=WD_ALIGN_PARAGRAPH.LEFT if j == 0
                             else WD_ALIGN_PARAGRAPH.CENTER)
            if widths:
                for j, w in enumerate(widths):
                    for row in table.rows:
                        row.cells[j].width = Cm(w)
            para(doc, '', space_after=8)
        elif kind == 'listing':
            _, caption, code = block
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.rows[0].cells[0]
            set_cell(cell, code, size=9.5, mono=True)
            shade_cell(cell, 'F5F5F0')
            cap = para(doc, resolve_refs(caption), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
                       space_before=4)
        elif kind == 'pagebreak':
            doc.add_page_break()
        else:
            raise ValueError(f'Nepoznat blok: {kind}')


def render_kdi(doc, title, rows):
    para(doc, title, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    for r in doc.paragraphs[-1].runs:
        r.bold = True
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        set_cell(table.rows[i].cells[0], k, size=10, bold=True)
        set_cell(table.rows[i].cells[1], v, size=10)
        table.rows[i].cells[0].width = Cm(6.2)
        table.rows[i].cells[1].width = Cm(10.3)


# ---------------------------------------------------------------- sklapanje

def main():
    doc = Document()

    # osnovni stilovi
    normal = doc.styles['Normal']
    set_style_font(normal, FONT, 12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    set_style_font(doc.styles['Heading 1'], FONT, 16, bold=True,
                   color=RGBColor(0, 0, 0))
    set_style_font(doc.styles['Heading 2'], FONT, 14, bold=True,
                   color=RGBColor(0, 0, 0))
    set_style_font(doc.styles['Heading 3'], FONT, 12.5, bold=True,
                   color=RGBColor(0, 0, 0))
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(0)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(14)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(14)
    doc.styles['Heading 2'].paragraph_format.space_after = Pt(8)
    doc.styles['Heading 3'].paragraph_format.space_before = Pt(10)
    doc.styles['Heading 3'].paragraph_format.space_after = Pt(6)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    add_page_number_footer(section)

    # ---------------- naslovna strana
    para(doc, '', space_after=0)
    para(doc, 'УНИВЕРЗИТЕТ У НОВОМ САДУ', size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    for r in doc.paragraphs[-1].runs:
        r.bold = True
    para(doc, 'ФАКУЛТЕТ ТЕХНИЧКИХ НАУКА У НОВОМ САДУ', size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for r in doc.paragraphs[-1].runs:
        r.bold = True
    for _ in range(7):
        para(doc, '', space_after=0)
    para(doc, 'Урош Петрашковић', size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=18)
    para(doc,
         'АУТОМАТСКО ГЕНЕРИСАЊЕ ПИТАЊА ЗА ПРОВЕРУ ЗНАЊА ПО СОЛО '
         'ТАКСОНОМИЈИ ПРИМЕНОМ ВЕЛИКИХ ЈЕЗИЧКИХ МОДЕЛА',
         size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    for r in doc.paragraphs[-1].runs:
        r.bold = True
    para(doc, 'МАСТЕР РАД', size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=2)
    for r in doc.paragraphs[-1].runs:
        r.bold = True
    para(doc, '__________________________', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, 'Мастер академске студије', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(8):
        para(doc, '', space_after=0)
    para(doc, 'Ментор: проф. др Горан Савић', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, 'Кандидат: Урош Петрашковић', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    for _ in range(3):
        para(doc, '', space_after=0)
    para(doc, 'Нови Сад, 2026.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---------------- KDI (srpski + engleski)
    doc.add_page_break()
    render_kdi(doc, 'КЉУЧНА ДОКУМЕНТАЦИЈСКА ИНФОРМАЦИЈА', KDI_SR)
    doc.add_page_break()
    render_kdi(doc, 'KEY WORDS DOCUMENTATION', KDI_EN)

    # ---------------- sadržaj
    doc.add_page_break()
    h = doc.add_heading('Садржај', level=1)
    for r in h.runs:
        r.font.name = FONT
    add_toc_field(doc)

    # ---------------- telo rada
    render_blocks(doc, BLOCKS_1)
    render_blocks(doc, BLOCKS_2)
    render_blocks(doc, BLOCKS_3)

    # ---------------- literatura (po redosledu pojavljivanja)
    doc.add_page_break()
    h = doc.add_heading('Литература', level=1)
    for r in h.runs:
        r.font.name = FONT
    for i, key in enumerate(ref_order, start=1):
        p = para(doc, f'[{i}] {REFS[key]}', size=11, space_after=5)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
    unused = [k for k in REFS if k not in ref_order]
    if unused:
        print('UPOZORENJE: reference bez pojavljivanja u tekstu:', unused)

    # ---------------- dodaci, pa biografija (uvek poslednja)
    render_blocks(doc, BLOCKS_DODACI)
    render_blocks(doc, BLOCKS_BIOGRAFIJA)

    doc.save(OUT_PATH)
    print('Sacuvano:', OUT_PATH)
    print(f'Referenci u literaturi: {len(ref_order)}')


if __name__ == '__main__':
    main()
