# -*- coding: utf-8 -*-
"""Generator master rada (MasterRad.docx). Sav tekst je na srpskom."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG = os.path.join(ROOT, 'comparison image.png')

doc = Document()

# --- osnovni stil ---
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


def P(text, align='justify', italic=False, size=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = 'Times New Roman'
    return h


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return t


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)


_FIGN = [0]
FIGDIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'figures')


def add_figure(path, cap_text, width=5.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _FIGN[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Slika {_FIGN[0]}. {cap_text}')
    r.italic = True
    r.font.size = Pt(10)


def FIG(fname, cap_text, width=5.8):
    add_figure(os.path.join(FIGDIR, fname), cap_text, width)


# ============================ NASLOVNA ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Univerzitet u Novom Sadu\nFakultet tehničkih nauka')
r.bold = True
r.font.size = Pt(14)
doc.add_paragraph()
doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run('Automatsko generisanje pitanja za proveru znanja po SOLO taksonomiji '
               'primenom velikih jezičkih modela')
r.bold = True
r.font.size = Pt(18)
doc.add_paragraph()
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
st.add_run('master rad').italic = True
for _ in range(6):
    doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run('Autor: Uroš Petrašković\nMentor: prof. dr Goran Savić').font.size = Pt(13)
doc.add_paragraph()
pl = doc.add_paragraph()
pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
pl.add_run('Novi Sad, 2026.').font.size = Pt(12)
doc.add_page_break()

# ============================ SAŽETAK ============================
H('Sažetak', 1)
P('Ručno sastavljanje kvalitetnih pitanja za proveru znanja je posao koji oduzima dosta '
  'vremena i traži od nastavnika i dobro poznavanje gradiva i poznavanje pravila za pisanje '
  'pitanja. Posebno je teško napraviti pitanja koja ciljaju različite nivoe razumevanja, jer '
  'tu nije dovoljno samo postaviti tačno pitanje, nego ga i smestiti na pravi kognitivni nivo. '
  'U ovom radu je opisan sistem koji taj posao automatizuje. Sistem iz nastavnog materijala u '
  'PDF formatu izdvaja strukturu kursa, gradi ontologiju pojmova i na osnovu nje velikim '
  'jezičkim modelom generiše pitanja višestrukog izbora raspoređena po SOLO taksonomiji, od '
  'unistrukturalnog do prošireno apstraktnog nivoa.')
P('Težište rada je na sloju za kontrolu kvaliteta. Pošto jezički model sam po sebi pravi dosta '
  'grešaka, oko generatora je izgrađen skup provera zasnovanih na pedagoškoj i psihometrijskoj '
  'literaturi, koje mere usklađenost pitanja sa zadatim SOLO nivoom, pokrivenost gradiva, '
  'tačnost odgovora, kvalitet distraktora, dvosmislenost i čitljivost. Kvalitet sistema je '
  'izmeren na slučaju kursa Operativni sistemi, a posebno je analiziran prelazak sa lokalnog '
  'modela na model dostupan preko servisa u oblaku, koji je doneo izrazito poboljšanje. Sistem '
  'je dodatno proveren i na referentnom skupu EduQG, koji sadrži pitanja koja su pisali '
  'stručnjaci, čime su kalibrisani sami merni instrumenti.')
P('Ključne reči: automatsko generisanje pitanja, SOLO taksonomija, ontologija, veliki jezički '
  'modeli, kvalitet pitanja višestrukog izbora.', italic=True)
doc.add_page_break()

# ============================ 1. UVOD ============================
H('1. Uvod', 1)
P('Obrazovanje je oduvek bilo jedan od temelja društva, a kako tehnologija postaje sve '
  'prisutnija u svakodnevnom životu, raste i potreba za alatima koji olakšavaju rad nastavnika '
  'i učenje studenata. Jedan od poslova koji nastavnicima oduzima najviše vremena jeste '
  'sastavljanje pitanja za proveru znanja. Da bi jedno pitanje bilo dobro, ono mora da bude '
  'tačno, jasno postavljeno, da pokriva važan deo gradiva i da ima uverljive netačne odgovore, '
  'a uz sve to treba da gađa određeni nivo razumevanja. Pisanje takvih pitanja u većem broju '
  'je naporno, pa se u praksi često dešava da provere znanja ostanu na nivou pukog prepoznavanja '
  'činjenica, dok dublje razumevanje ostane neprovereno.')
P('Ovaj rad se bavi tim problemom kroz sistem koji nastavni materijal, konkretno predavanja u '
  'PDF formatu, pretvara u pitanja za proveru znanja. Za razliku od opštih generatora pitanja, '
  'ovaj sistem pitanja razvrstava po SOLO taksonomiji (Structure of Observed Learning Outcomes). '
  'SOLO taksonomija opisuje kako razumevanje neke teme postepeno raste u složenosti i omogućava '
  'da se pitanja prave na različitim kognitivnim nivoima, od onih koja proveravaju jednu '
  'činjenicu do onih koja traže primenu naučenog u novom kontekstu. Na taj način se otvara '
  'mogućnost da se provere prilagode tome koliko duboko student vlada gradivom.')
P('Da bi generisanje bilo vezano za samo gradivo, a ne za opšte znanje modela, sistem prvo iz '
  'materijala izdvaja strukturu kursa. Ta struktura je hijerarhijska: kurs sadrži lekcije, '
  'lekcije sadrže sekcije, a sekcije sadrže pojedinačne nastavne objekte. Veze između nastavnih '
  'objekata se čuvaju u obliku ontologije, koja formalno opisuje kako su delovi gradiva povezani. '
  'Ontologija kasnije služi i kao osnova za pitanja viših nivoa, gde je bitno da pitanje testira '
  'tačno određenu vezu između pojmova, a ne neku proizvoljnu vezu koju bi model sam izmislio.')
P('Početna verzija ovog sistema je opisana u kraćem konferencijskom radu, gde je kvalitet '
  'pitanja procenjen kvalitativno, pregledom po jednog primera za svaki SOLO nivo. Taj rad je '
  'pokazao osnovnu ideju, ali je ostavio otvoreno pitanje kako kvalitet izmeriti na sistematičan '
  'način i kako ga poboljšati. Master rad nastavlja tu priču na tri načina. Prvo, oko generatora '
  'je izgrađen sloj za kontrolu kvaliteta, sastavljen od dvanaest provera koje su zasnovane na '
  'literaturi iz oblasti pisanja test pitanja i psihometrije. Drugo, kvalitet je izmeren '
  'brojkama, po lekcijama i zbirno, i to pre i posle prelaska na jači jezički model. Treće, '
  'sistem je proveren i na spoljnom referentnom skupu pitanja koja su pisali stručnjaci, čime '
  'su kalibrisani sami merni instrumenti.')
P('Procena kvaliteta je organizovana oko četiri istraživačka pitanja:')
bullet('IP1: Da li generisana pitanja zaista odgovaraju SOLO nivou koji im je dodeljen?')
bullet('IP2: Da li su distraktori, odnosno netačni odgovori, uverljivi, međusobno različiti i '
       'jasno netačni, a ne parafraze tačnog odgovora?')
bullet('IP3: Koliko dobro generisana pitanja pokrivaju nastavni materijal?')
bullet('IP4: Da li su pitanja tačna, korektno postavljena i bez halucinacija, odnosno da li je '
       'tačan odgovor zaista potkrepljen materijalom?')
P('Svako od ovih pitanja je u radu povezano sa konkretnom merom, pa se na njih ne odgovara '
  'utiskom, nego brojem. U nastavku se prvo daje pregled srodnih istraživanja i pokazuje koji '
  'je deo svakog rada primenjen u ovom rešenju, zatim se opisuje sam sistem i njegova '
  'implementacija, pa se iznose izmereni rezultati, opisuje prototip aplikacije i na kraju '
  'izvode zaključci.')

# ============================ 2. SRODNA ISTRAZIVANJA ============================
H('2. Srodna istraživanja', 1)
P('Oblast automatskog generisanja pitanja ima dugu istoriju, od ranih sistema zasnovanih na '
  'pravilima do današnjih pristupa koji se oslanjaju na velike jezičke modele. Jedan od '
  'začetnika oblasti je rad Mitkova (2003) [3], koji je koristio pravila obrade prirodnog '
  'jezika da izjavne rečenice pretvori u pitanja. Takav pristup je dobro radio za jednostavne '
  'činjenice, na primer rečenicu o tome gde se odvija fotosinteza je pretvarao u pitanje gde se '
  'fotosinteza odvija, ali se mučio sa složenijim gradivom i često je davao gramatički nezgrapna '
  'pitanja. Kasneci i saradnici (2023) [4] dali su pregled primene jezičkih modela poput '
  'ChatGPT-a u obrazovanju i ukazali i na potencijal i na rizike, pre svega na pitanja tačnosti, '
  'pristrasnosti i potrebe za nadzorom nastavnika. Lister i saradnici (2006) [5] pokazali su '
  'kako se SOLO taksonomija može primeniti za klasifikaciju programerskih pitanja, od '
  'unistrukturalnih, kao što je pitanje šta jedna promenljiva čuva, do prošireno apstraktnih, '
  'kao što je zahtev da se osmisli program koji rešava složen problem, i upozorili da se provere '
  'previše oslanjaju na niže nivoe znanja, a zapostavljaju više. Liang i saradnici (2018) [6] '
  'bavili su se generisanjem distraktora pristupom rangiranja, gde se napravi mnogo kandidata pa '
  'se biraju oni najuverljiviji, uz zapažanje da dobar distraktor mora biti blizak tačnom '
  'odgovoru po značenju, ali pažljivim čitanjem jasno netačan; ako je tačan odgovor mitohondrija, '
  'dobri distraktori su ribozomi ili jedro, a ne nepovezani pojmovi.')
P('Sistem opisan u ovom radu nadovezuje se na ta istraživanja, ali ide korak dalje tako što oko '
  'generatora postavlja sloj provera. U nastavku se navode radovi koji su konkretno primenjeni, '
  'uz objašnjenje kog dela rešenja se svaki tiče.')

H('2.1. Generisanje pitanja i navođenje modela', 2)
P('Polazna tačka za generisanje pitanja je rad Scaria i saradnika (2024) [7], koji se bavi '
  'generisanjem pitanja po Blumovoj taksonomiji uz pomoć jezičkih modela. Iako taj rad koristi '
  'Blumovu, a ne SOLO taksonomiju, iz njega su preuzeti vredni zaključci o tome kako oblikovati '
  'prompt. Prvo, prompt počinje uputstvom u kojem se modelu zadaje uloga stručnjaka za '
  'sastavljanje provera znanja, jer takvo uokvirivanje uloge primetno poboljšava rezultat. '
  'Drugo, traženi kognitivni nivo se opisuje sa svega jednom ili dve rečenice, da definicija '
  'bude jasna a ne preopširna. Treće, modelu se daje primer dobrog pitanja, ali ne previše '
  'primera; u tom radu se pokazalo da strategija sa pet primera radi lošije, pa je u ovom '
  'sistemu zadržan samo jedan primer. Četvrto, koristi se takozvana lestvica za razmišljanje '
  '(chain of thought), gde se od modela traži da razmišlja korak po korak pre nego što da '
  'konačan odgovor. Bez toga model često sve odradi u jednom potezu, bez provere pojedinačnih '
  'koraka, što vodi do haotičnih pitanja, izmišljenih pojmova i loših distraktora. Peto, uz '
  'svako pitanje se čuva i doslovan citat iz izvornog materijala koji opravdava tačan odgovor; '
  'ako takav citat ne postoji, to je znak da je odgovor verovatno izmišljen.')
P('Lestvica za razmišljanje detaljnije je obrađena u radu Wei i saradnika (2022) [8], koji je u '
  'celini posvećen tehnici chain-of-thought navođenja. To je zapravo ista ideja kao peta tačka '
  'prethodnog rada, samo razrađena. Pravljenje pitanja samo jednim pozivom modela nema dovoljno '
  'strukture, rezultati dosta variraju i ima previše grešaka. Ako se posao razbije na korake i '
  'model se drži preciznih međukoraka, rezultati su znatno bolji, naročito na zadacima koji '
  'zahtevaju zaključivanje. U ovom sistemu se taj princip koristi i pri generisanju i pri '
  'kasnijoj proveri pitanja.')
P('Pitanje koliko primera dati i kakvi treba da budu povezano je sa teorijom kognitivnog '
  'opterećenja, koju su postavili Sweller i Cooper (1985) [9]. Oni su pokazali da rešeni primeri '
  'znatno pomažu učenju, što se prenosi i na navođenje modela. Iz tog razloga se u prompt '
  'ubacuje primer dobrog pitanja, ali je važan jedan detalj: primer nije iz oblasti koja se '
  'trenutno obrađuje. Drugim rečima, ako se generišu pitanja iz operativnih sistema, primer u '
  'promptu je iz sasvim druge teme. Tako model preuzima samo formu dobrog pitanja, a ne i temu. '
  'U ovom sistemu je za primer uzeta fotosinteza, koja nema nikakve veze sa gradivom kurseva, '
  'pa nema opasnosti da model prepiše sadržaj umesto strukture.')

H('2.2. Borba protiv halucinacija', 2)
P('Veliki problem kod jezičkih modela je halucinacija, odnosno izmišljanje sadržaja kojeg nema '
  'u materijalu. Prvi mehanizam zaštite oslanja se na ideju iz rada Lewis i saradnika (2020) '
  '[10] o generisanju potpomognutom pretragom (retrieval-augmented generation). U tom pristupu '
  'se odgovor vezuje za izvor. U ovom sistemu se to ostvaruje poljem koje se zove source_line, '
  'a to je doslovan citat iz PDF materijala koji opravdava tačan odgovor. Ako se taj citat ne '
  'može pronaći u materijalu, to je jak znak da je odgovor izmišljen. Polje se čuva i prikazuje, '
  'pa nastavnik koji pregleda pitanja može da uporedi citat sa odgovorom i, ako se ne poklapaju, '
  'da pitanje ispravi ili obriše. Bitno je naglasiti da se ova provera odnosi pre svega na tačan '
  'odgovor, jer je najvažnije da baš on ne bude halucinacija.')
P('Source_line rešava slučaj kada model izmisli nešto čega nema u materijalu, ali ne rešava '
  'suptilniji problem, a to je kada citat postoji, ali je pogrešno protumačen. Za taj slučaj se '
  'koristi metoda iz rada Dhuliawala i saradnika (2024) [15] o lancu provere (chain of '
  'verification). Umesto da se model jednostavno pita da li je odgovor tačan, na šta bi verovatno '
  'lažno potvrdio, postupak ide u četiri koraka. Najpre se uzme pitanje i tačan odgovor. Zatim '
  'se isplaniraju dva do tri kratka pomoćna pitanja čiji odgovori, zajedno gledani, potvrđuju '
  'ili obaraju tačan odgovor. Potom se na svako pomoćno pitanje odgovara nezavisno, koristeći '
  'samo izvorni materijal. Na kraju se donosi sud da li je tačan odgovor potvrđen, neodređen ili '
  'oboren. Pitanja koja su neodređena ili oborena se označavaju za ljudski pregled.')

H('2.3. Generisanje distraktora', 2)
P('Način generisanja distraktora za najsloženija pitanja preuzet je iz rada Bitew i saradnika '
  '(2023) [11] o generisanju distraktora prediktivnim navođenjem. Pošto su prošireno apstraktna '
  'pitanja najteža, kod njih se distraktori prave u dva prolaza. U prvom prolazu model osmisli i '
  'zapamti tačan odgovor, objašnjenje, citat iz materijala i sam tekst pitanja. Drugi prolaz '
  'služi isključivo za pravljenje tri distraktora, gde se model trudi da ih napiše kao stvarne '
  'zablude i da svaki sledi drugačiju strategiju. Veoma važna pouka iz tog rada je da se pazi da '
  'distraktori ne budu slični i da dva distraktora ne predstavljaju isti pogrešan koncept samo '
  'drugačije sročen. Pošto model ima ograničenu pažnju, kvalitet distraktora je primetno bolji '
  'kada se oni ne prave u isto vreme kada i sve ostalo, nego se za njih napravi poseban poziv. '
  'U pomenutom radu se navodi da je za prošireno apstraktna pitanja oko 53 procenta distraktora '
  'spremno za upotrebu, što govori koliko je taj nivo težak.')
P('Da bi posao bio sistematičan, za svaki SOLO nivo su definisani tipovi distraktora, čime '
  'generisanje postaje gotovo šablonsko:')
bullet('Unistrukturalni: VARIANT_OF_CORRECT, NEAR_SYNONYM, COMMON_MISCONCEPTION.')
bullet('Multistrukturalni: LIST_WITH_ONE_WRONG_ITEM, LIST_MISSING_ONE_ITEM, '
       'RELATED_BUT_OUT_OF_SCOPE.')
bullet('Relacioni: REVERSED_CAUSE_EFFECT, CORRELATION_AS_CAUSATION, '
       'DIFFERENT_REAL_RELATIONSHIP.')
bullet('Prošireno apstraktni: APPLIES_WRONG_PRINCIPLE, RIGHT_PRINCIPLE_WRONG_DOMAIN, '
       'OVER_GENERALIZATION.')
P('Drugi važan deo ovog rada je upotreba embedinga, gde se tekst pretvara u vektore, pri čemu '
  'sličan tekst ima sličan vektor. Tu se rad naslanja na BERTScore, koji su predložili Zhang i '
  'saradnici (2019) [12]. Koristi se mera kosinusne sličnosti, koja govori koliko su dva vektora '
  'slična po smeru. Rezultat je broj između minus jedan i jedan, gde jedan znači gotovo isto '
  'značenje, oko 0,9 vrlo slično, oko 0,5 blisko ali ipak različito, nula nepovezano, a '
  'negativne vrednosti suprotno značenje. U kodu postoje pragovi za najmanju i najveću '
  'uverljivost: distraktor čija je sličnost sa tačnim odgovorom veća od 0,92 smatra se '
  'parafrazom, a onaj sa sličnošću manjom od 0,40 smatra se previše očigledno netačnim. U oba '
  'slučaja se postavlja upozorenje vredno pet poena, gde veći broj poena znači lošije pitanje, '
  'i modelu se sugeriše da taj distraktor zameni boljim. Ti pragovi nisu proizvoljni, nego su '
  'dobijeni kalibracijom. Na isti način se distraktori porede i međusobno, jer je čest problem '
  'bio da dva distraktora znače istu stvar; uvođenje embedinga i pravila iz literature znatno je '
  'popravilo tu stranu generisanja.')

H('2.4. Pitanja relacionog nivoa i ontologija', 2)
P('Za generisanje pitanja relacionog nivoa korišćen je rad Chen i Shiu (2025) [13] o sistemu '
  'KAQG, koji spaja graf znanja i generisanje potpomognuto pretragom radi kontrole težine '
  'pitanja. Glavna pouka je da je neophodno koristiti sidro, odnosno tačno zadatu vezu u grafu '
  'znanja. Ako se modelu ne kaže kojom se vezom tačno bavi, on počinje da luta i pitanja '
  'skreću ka uopštenim poređenjima. Sidro tu opštost sprečava i usmerava model na tačno '
  'određeni deo grafa koji je izgrađen iz ontologije. Ako, na primer, postoji veza tipa '
  'preduslov, modelu se izričito kaže da mora da testira baš tu vezu, a ne neku drugu koju bi '
  'mogao da pretpostavi. U sklopu ovog dela je iznova napisan ceo postupak generisanja '
  'ontologije, sa težištem na kvalitetu, pa je rezultat manji broj veza koje su zato znatno '
  'pouzdanije. Sve veze se čuvaju i kasnije porede sa tim koliko ih pitanja pokrivaju.')

H('2.5. Pravila za pisanje pitanja i provera kvaliteta', 2)
P('Osnovni skup pravila za pisanje pitanja višestrukog izbora preuzet je iz rada Haladyna, '
  'Downing i Rodriguez (2002) [14], koji sadrži čak trideset jedno pravilo, proverena kroz '
  'dvadeset sedam empirijskih studija i veliki broj udžbenika. U sistemu se koriste dve grupe '
  'pravila: STEM_RULES za tekst pitanja i OPTION_RULES za odgovore. Činjenica da su pravila '
  'numerisana ima praktičan značaj, jer ih model tako mnogo bolje prati nego kada su data kao '
  'slobodan tekst. Pravila se najpre stavljaju u prompt kao mera prevencije, a zatim se naknadno '
  'proverava da li ih se model držao. Neka od pravila su:')
bullet('Pravilo teksta 1: pitanje se završava upitnikom ili jasnim imperativom.')
bullet('Pravilo teksta 2: tekst je kraći od oko 250 znakova, bez nepotrebnog ukrasa.')
bullet('Pravilo teksta 3: izbegavati negaciju, a ako se koristi, istaći je.')
bullet('Pravilo opcija 1: tačno jedna opcija je tačna.')
bullet('Pravilo opcija 2: nijedne dve opcije nisu parafraze jedna druge.')
bullet('Pravilo opcija 3: najduža opcija nije duža od dvostruke najkraće.')
bullet('Pravilo opcija 4: tačna opcija ne sme biti najduža.')
bullet('Pravilo opcija 5: ne koristiti odgovore tipa svi navedeni ili nijedan navedeni.')
bullet('Pravilo opcija 6: numeričke opcije idu rastuće ili opadajuće.')
bullet('Pravilo opcija 7: opcije imaju istu gramatičku strukturu, istu vrstu reči i isti '
       'register.')
P('Ovaj rad je posebno pomogao kod distraktora. Ako je tačan odgovor glagol, a distraktori '
  'imenice ili obrnuto, to studentu olakšava da pogodi tačan odgovor, na šta se odnosi pravilo '
  'opcija 7. Slično, ako je tačan odgovor osetno duži od distraktora, to je trag, što pokrivaju '
  'pravila opcija 3 i 4. Rad uvodi i sloj provere kvaliteta i valjanosti pre upotrebe, gde '
  'pitanje nakon generisanja prolazi kroz validaciju. U ovom sistemu je taj princip ostvaren '
  'tako što jedanaest pravila proverava ne poziv modela, nego običan programski mehanizam za '
  'pravila, koji skenira pitanje. Mehanizam razlikuje grešku, koja vredi petnaest poena, i '
  'upozorenje, koje vredi pet poena, pa se na kraju prikazuje koliko je grešaka i upozorenja '
  'pokrenuto, čime se dobija ocena kvaliteta po ovoj meri. Treba naglasiti da se za proveru '
  'pravila koriste oba pristupa: u promptu koji radi prevenciju koristi se jezički model, a u '
  'drugom prolazu se ne koristi model, nego samo mehanizam za pravila proverava da li su pravila '
  'ispoštovana.')
P('Posebna pažnja se poklanja pravilu da tekst pitanja mora da nosi glavnu misao, tako da bi '
  'onaj ko rešava test trebalo da može da odgovori već iz teksta pitanja, pre nego što uopšte '
  'pogleda ponuđene opcije. Ovo pravilo se ne može proveriti na uobičajen način, nego se mora '
  'pozvati model. Model pokušava da odgovori na pitanje bez ponuđenih opcija, a zatim se taj '
  'odgovor poredi sa tačnim odgovorom da bi se videlo koliko se poklapaju. To je dodatni prolaz '
  'modela koji daje koeficijent poklapanja i znatno produžava vreme obrade, ali se isplati jer '
  'otkriva pitanja koja se ne mogu rešiti bez gledanja u opcije.')

H('2.6. Procena težine i klasifikacija nivoa', 2)
P('Za procenu težine pitanja korišćen je pristup iz klasične teorije testova, koju su izložili '
  'Crocker i Algina (1986) [16]. Tačan odgovor se sakrije od modela, pa se model pita da reši '
  'pitanje, i to se ponovi određeni broj puta da bi se dobila mera koliko dobro pogađa. Uvodi '
  'se koeficijent p koji govori koliko je pitanje lako. Ako je p jednako jedan, pitanje je '
  'trivijalno ili postoji neki trag koji ga odaje, na primer dužina ili gramatička razlika '
  'opcija. Ako je p između 0,6 i 0,9, reč je o primerenoj težini za model koji rešava. Ako je p '
  'manje od 0,5, pitanje je ili loše postavljeno, ili je odgovor pogrešan, ili se odgovor ne '
  'nalazi u materijalu. Ovaj pristup automatski pronalazi neadekvatna pitanja, pa iako su to '
  'dodatni pozivi modela, oni se isplate jer filtriraju loša pitanja bez da na to čovek troši '
  'vreme.')
P('Da bi se proverilo da li pitanja zaista pripadaju nivou koji im je dodeljen, koristi se drugi '
  'jezički model koji ne zna koji je nivo tražen i koji svako pitanje samostalno klasifikuje. '
  'Zatim se meri koliko se njegova procena slaže sa zadatim SOLO nivoom. Slaganje se ne meri '
  'prostim procentom pogodaka, jer bi to bilo naivno: ako je većina pitanja unistrukturalna, '
  'pošto su ona najlakša za generisanje, model bi mogao da postigne visok procenat samo tako '
  'što uvek tipuje na najčešći nivo. Zato se koristi Koenov koeficijent kapa, predstavljen u '
  'radu Cohen (1960) [18], a tumačenje vrednosti prati skalu iz rada Landis i Koch (1977) [17]. '
  'Kapa u svoju formulu uračunava slaganje koje bi nastalo i pukim slučajem, pa daje realniju '
  'sliku. Vrednost jedan znači savršeno slaganje, vrednost između 0,6 i jedan znači da je '
  'generator dobar, nula znači slaganje na nivou slučajnosti, a negativna vrednost znači da je '
  'generator gori nego nasumično pogađanje. Uz kapa vrednost se prikazuje i matrica konfuzije, '
  'koja pokazuje na kom se nivou koliko greši.')

H('2.7. Pokrivenost gradiva, čitljivost, dvosmislenost i zablude', 2)
P('Za procenu koliko pitanja pokrivaju gradivo oslonac je rad Kurdi i saradnika (2020) [19], '
  'sistematski pregled automatskog generisanja pitanja u obrazovne svrhe. Starija verzija '
  'sistema merila je pokrivenost brojem stranica i znakova, što ne govori ništa o semantičkoj '
  'pokrivenosti. Zato se ovde koristi pokrivenost pojmova, gde je skup pojmova unija nastavnih '
  'objekata i svih njihovih ključnih reči. Taj rad je uveo pokrivenost zasnovanu na grafu '
  'znanja i skrenuo pažnju na slabosti merenja preko stranica. U ovom sistemu se za svaki pojam '
  'gleda koliko veza ima u ontološkom grafu, pa pojam koji se češće pominje ima veću težinu. '
  'Pojam kao što je proces je centralan i povezuje sve ostalo, pa ako ga pitanje pominje, '
  'pokrilo je važan deo gradiva, dok ređi pojmovi nose manju težinu. Kod višerečnih pojmova '
  'reči se filtriraju tako da moraju biti jedna za drugom, jer ako se dve reči nađu u istoj '
  'rečenici, to ne mora da znači da je reč o tom pojmu. Na kraju se prikazuju ponderisana i '
  'neponderisana pokrivenost, gde ponderisana uzima u obzir koliko je važan pojam koji je '
  'pokriven.')
P('Čitljivost pitanja meri se formulama iz radova Flesch (1948) [20] i Kincaid i saradnici '
  '(1975) [21]. Te formule koriste prosečan broj reči po rečenici i prosečan broj slogova po '
  'reči i procenjuju koji bi školski razred mogao sa razumevanjem da pročita tekst. Cilj je da '
  'pitanja nižih kognitivnih nivoa imaju manji zahtev za čitanje, odnosno da budu lakša za '
  'čitanje. Ova provera ne koristi poziv modela.')
P('Dvosmislenost se procenjuje na osnovu rada Downing (2005) [22] o posledicama kršenja pravila '
  'za pisanje pitanja. Reč je o slučaju kada se nešto sroči tako da može imati dva značenja, pa '
  'onaj ko rešava test razume pitanje drugačije nego što je pisac mislio. Ono što je nastavniku '
  'odmah jasno može studenta da navede na pogrešan trag ako su pojmovi dvosmisleni. Ovde se '
  'koristi model, a prompt izričito traži da model izdvoji alternativna tumačenja, a ne da samo '
  'kaže da li je nešto nejasno. Kada se otkrije dvosmislen pojam, dodaje se pet poena upozorenja '
  'i beleži se razlog, koji korisnik kasnije može da vidi.')
P('Najzad, najbolji distraktori se prave po uzoru na rad Sadler (1998) [23], koji pokazuje da '
  'se najdiskriminativniji distraktori dobijaju praćenjem samog materijala, jer u materijalu '
  'često već piše šta se obično brka i gde se greši. Umesto da model sam izmišlja zablude, one '
  'se vade iz materijala. Koriste se obrasci kao što su česta greška, studenti često misle, za '
  'razliku od, ne treba mešati, a common error, students often think, unlike X, Y is i not to '
  'be confused with. Pronađeni delovi se podele na odlomke od oko dvesta znakova i šalju modelu, '
  'koji procenjuje da li je reč o stvarnoj zabludi ili o lažnom pogotku obrasca.')

H('2.8. Referentni skup za evaluaciju', 2)
P('Za spoljnu proveru sistema korišćen je skup EduQG, koji su predstavili Hadifar i saradnici '
  '(2023) [24]. To je skup od 3397 pitanja višestrukog izbora koja su pisali stručnjaci, '
  'izvučenih iz dvanaest udžbenika, gde je svako pitanje vezano za izvorni tekst na nivou '
  'rečenice. Pošto su ta pitanja pisali stručnjaci, ona služe kao spoljni zlatni standard. '
  'Time se ne proverava samo generator, nego i sami merni instrumenti: ako provere previše '
  'često označavaju stručnjačka pitanja kao loša, znači da su loše kalibrisane. Ova uloga '
  'EduQG skupa detaljnije je opisana u poglavlju o rezultatima.')

# ============================ 3. SISTEM ============================
H('3. Sistem za kreiranje pitanja po SOLO taksonomiji', 1)
H('3.1. SOLO taksonomija', 2)
P('SOLO taksonomija opisuje kako razumevanje neke teme raste u složenosti i sastoji se od pet '
  'nivoa. Na prestrukturalnom nivou student još nema razumevanje teme i odgovori su uglavnom '
  'nebitni ili izostaju. Na unistrukturalnom nivou razume jedan aspekt. Na multistrukturalnom '
  'razume više aspekata, ali ne vidi kako su povezani. Na relacionom nivou ume da poveže delove '
  'u celinu. Na prošireno apstraktnom nivou ume da uopšti i primeni znanje u novim situacijama. '
  'Sistem generiše pitanja za četiri viša nivoa, jer prestrukturalni nivo po definiciji ne '
  'odgovara nijednom smislenom pitanju. Unistrukturalna pitanja se tiču jedne činjenice, '
  'multistrukturalna više činjenica, relaciona traže povezivanje pojmova, a prošireno '
  'apstraktna traže primenu naučenog u novom kontekstu.')

FIG('fig8_solo_ladder.png',
    'SOLO taksonomija: četiri nivoa koja sistem generiše, poređana po porastu kognitivne '
    'složenosti.', 5.6)

H('3.2. Arhitektura i tok obrade', 2)
P('Sistem je izveden kao veb aplikacija sa troslojnom monolitnom arhitekturom. Korisnički deo '
  'je napravljen u tehnologiji React, serverski deo u okviru Flask na jeziku Python, a podaci '
  'se čuvaju u bazi SQLite. Pored relacione baze, struktura kursa se čuva i u obliku ontologije, '
  'nad kojom se mogu postavljati upiti jezikom SPARQL. Osnovna funkcija sistema je automatsko '
  'pravljenje provera znanja po SOLO taksonomiji, kroz tok obrade koji sirov materijal pretvara '
  'u struktuirane kvizove.')
P('Tok počinje učitavanjem nastavnog materijala u PDF formatu, nakon čega se tekst šalje modelu '
  'na raščlanjivanje. Arhitektura ne zavisi od konkretnog modela, pa se može koristiti i lokalni '
  'i model u oblaku. Model iz sirovog materijala izdvaja strukturu kursa i smešta je u bazu kao '
  'hijerarhiju: kurs sadrži lekcije, lekcije sadrže sekcije, a sekcije sadrže nastavne objekte. '
  'Kada su nastavni objekti napravljeni, sistem automatski uspostavlja ontološke veze između '
  'njih, koje se mogu izvesti i kao OWL datoteka. Zatim sistem prolazi kroz nastavne objekte i '
  'veze i generiše pitanja za sva četiri SOLO nivoa, pri čemu se za svaki nivo koristi prompt '
  'prilagođen traženoj složenosti. Svako pitanje se čuva zajedno sa opcijama, tačnim odgovorom, '
  'objašnjenjem i citatom iz materijala. Na kraju nastavnik bira teme i broj pitanja po nivou, '
  'pa sistem od tog skupa sastavlja proveru.')
FIG('fig7_pipeline.png',
    'Tok obrade sistema, od nastavnog materijala u PDF formatu do gotove provere znanja.', 6.4)

H('3.3. Generisanje pitanja po nivoima', 2)
P('Svaki nivo se generiše posebnim promptom. Zajednički delovi svih promptova su uloga '
  'stručnjaka za sastavljanje provera, kratak opis traženog nivoa, jedan primer dobrog pitanja '
  'iz nepovezane oblasti, lestvica za razmišljanje i pravila za tekst i opcije. Unistrukturalna '
  'i multistrukturalna pitanja se generišu u jednom prolazu. Relaciona pitanja se vezuju za '
  'tačno određenu vezu iz ontologije, koja služi kao sidro, da pitanje ne bi skliznulo u opšte '
  'poređenje. Prošireno apstraktna pitanja se generišu u dva prolaza. U prvom prolazu se '
  'osmišljavaju tekst pitanja, tačan odgovor, objašnjenje i citat, a u drugom se prave '
  'distraktori, svaki po svojoj strategiji. Razlog za dva prolaza je ograničena pažnja modela: '
  'kada istovremeno mora da smisli i scenario i distraktore, kvalitet opada, pa se poseban '
  'poziv za distraktore isplati.')

H('3.4. Sloj za kontrolu kvaliteta', 2)
P('Najveća razlika u odnosu na početnu verziju sistema je sloj za kontrolu kvaliteta. On radi '
  'na dva načina. Prvi je prevencija, gde se pravila i smernice ugrađuju u sam prompt, pa model '
  'pokušava da grešku ne napravi. Drugi je naknadna provera, gde se već generisano pitanje '
  'proverava, delom običnim programskim mehanizmom za pravila, a delom dodatnim pozivima modela. '
  'Ukupno se koristi dvanaest provera, koje pokrivaju četiri istraživačka pitanja iz uvoda. '
  'Tabela u nastavku daje pregled provera, izvora na kojem se zasnivaju i istraživačkog pitanja '
  'na koje se odnose.')
table(
    ['Provera', 'Izvor', 'Tiče se'],
    [
        ['Pokrivenost pojmova', 'Kurdi 2020 [19]', 'IP3'],
        ['Haladyna pravila (rule engine)', 'Haladyna 2002 [14]', 'IP2, IP4'],
        ['Klasifikacija SOLO nivoa (kapa)', 'Cohen 1960 [18], Landis i Koch 1977 [17]', 'IP1'],
        ['Lanac provere (CoVe)', 'Dhuliawala 2024 [15]', 'IP4'],
        ['Procena težine (p vrednost)', 'Crocker i Algina 1986 [16]', 'IP4'],
        ['Rešivost iz teksta (stem only)', 'Haladyna 2002 [14]', 'IP4'],
        ['Indeks podudaranja sa ciljem (IOC)', 'Haladyna 2002 [14]', 'IP1, IP4'],
        ['Čitljivost (Flesch, Kincaid)', 'Flesch 1948 [20], Kincaid 1975 [21]', 'IP1'],
        ['Dvosmislenost', 'Downing 2005 [22]', 'IP4'],
        ['Gramatička ujednačenost opcija', 'Haladyna 2002 [14]', 'IP2'],
        ['Uverljivost distraktora (face validity)', 'Liang 2018 [6], Zhang 2019 [12]', 'IP2'],
        ['Vađenje zabluda iz materijala', 'Sadler 1998 [23]', 'IP2'],
    ],
)
caption('Tabela 1. Provere kvaliteta, izvori i istraživačka pitanja na koja se odnose.')
P('Provere koje ne traže poziv modela, kao što su pravila pisanja, pokrivenost pojmova i '
  'čitljivost, su brze i deterministične. Provere koje traže model, kao što su klasifikacija '
  'nivoa, lanac provere, procena težine, rešivost iz teksta, dvosmislenost, gramatička '
  'ujednačenost i uverljivost, su sporije i čine da generisanje i provera traju osetno duže. '
  'Da se isti rezultat ne bi računao više puta, skupi rezultati se keširaju, kako pojedinačni '
  'pozivi modela, tako i celi izveštaji po lekciji, pa se pri ponovnom otvaranju lekcije '
  'rezultati odmah prikazuju.')

H('3.5. Lokalni i globalni model', 2)
P('Pošto sistem ne zavisi od konkretnog modela, isti tok obrade se može pokrenuti i sa lokalnim '
  'modelom i sa modelom u oblaku. U prvoj fazi rada korišćen je lokalni model Qwen 2.5 sa '
  'četrnaest milijardi parametara, koji se može pokrenuti na ličnom računaru bez interneta i '
  'bez ključeva za pristup. Kasnije je sistem prebačen na model Claude Haiku 4.5, dostupan preko '
  'servisa u oblaku. Kao što će biti pokazano u poglavlju o rezultatima, ta promena je donela '
  'izrazito poboljšanje, naročito kod provera koje traže zaključivanje. To je i očekivano, jer '
  'sloj za kontrolu kvaliteta u velikoj meri zavisi od sposobnosti modela da sudi, a tu lokalni '
  'model od četrnaest milijardi parametara nije bio dovoljan.')

# ============================ 4. REZULTATI ============================
H('4. Rezultati', 1)
H('4.1. Opis studije slučaja', 2)
P('Kvalitet sistema je izmeren na kursu Operativni sistemi, koji se sastoji od tri lekcije: '
  'Procesi, Niti i Konkurentnost. Iz tih lekcija je ukupno generisano 246 pitanja, i to 90 za '
  'lekciju Procesi, 66 za lekciju Niti i 90 za lekciju Konkurentnost. Po SOLO nivoima, '
  'najbrojnija su unistrukturalna pitanja, kojih ima 111, zatim multistrukturalna sa 75, '
  'relaciona sa 50 i prošireno apstraktna sa 10. Ovakav raspored je očekivan, jer su niža '
  'pitanja lakša za generisanje i ima ih više, što je upravo razlog zašto se slaganje nivoa '
  'meri Koenovim koeficijentom, a ne prostim procentom. Pored kvantitativnog merenja na kursu '
  'Operativni sistemi, u kvalitativnoj analizi koja sledi koriste se primeri i iz kursa '
  'Testiranje softvera, kao što je rađeno i u početnom radu.')
table(
    ['Lekcija', 'Broj pitanja', 'Uni', 'Multi', 'Relaciona', 'Proš. apstr.'],
    [
        ['Procesi', '90', '40', '27', '18', '5'],
        ['Niti', '66', '31', '18', '12', '5'],
        ['Konkurentnost', '90', '40', '30', '20', '0'],
        ['Ukupno', '246', '111', '75', '50', '10'],
    ],
)
caption('Tabela 2. Pregled studije slučaja i raspodela pitanja po SOLO nivoima.')
FIG('fig4_solo_distribution.png',
    'Raspodela generisanih pitanja po SOLO nivoima i lekcijama.', 5.6)

H('4.2. Kvalitativna analiza generisanih pitanja po nivoima', 2)
P('Pre brojki, korisno je pogledati po jedan tipičan primer za svaki SOLO nivo i prokomentarisati '
  'kvalitet pitanja i njegovih distraktora, kao što je rađeno u početnom radu. Primeri su uzeti '
  'iz kurseva Operativni sistemi i Testiranje softvera. Ova analiza je kvalitativna i služi kao '
  'dopuna mernim rezultatima koji slede.')
H('Unistrukturalni nivo', 3)
P('Unistrukturalna pitanja izdvajaju jedan osnovni pojam, što je glavni zahtev za ovaj nivo. Kao '
  'primer uzimamo pitanje: Programski brojač (PC) sadrži memorijsku adresu koje vrste entiteta? '
  'Tačan odgovor je sledeća instrukcija koja treba da se izvrši.')
P('Distraktori ovog pitanja su pokazali promenljiv kvalitet. Opcija trenutno izvršavana '
  'instrukcija je posebno uspešna, jer cilja vremensku zabunu između reči sledeća i trenutna, pa '
  'tačno pogađa studente koji nepotpuno razumeju ciklus pribavljanja i izvršavanja. Opcija sve '
  'instrukcije u memoriji cilja čestu početničku zabludu o opsegu registra, mada može biti '
  'očigledna studentima koji znaju da registar drži jednu vrednost. Opcija podaci i promenljive '
  'je najslabiji distraktor, jer studenti sa osnovnim znanjem arhitekture lako razlikuju adrese '
  'instrukcija i adrese podataka, pa je brzo odbacuju.')
H('Multistrukturalni nivo', 3)
P('Multistrukturalna pitanja traže da student barata sa više nezavisnih aspekata teme, obično '
  'kroz poređenje ili nabrajanje više tačnih osobina. Kao primer uzimamo pitanje: Koje od '
  'sledećih tvrdnji tačno opisuju i niti na korisničkom nivou (ULT) i niti na nivou jezgra (KLT)? '
  'Tačan odgovor je da ULT-ovima upravlja aplikacija, a KLT-ovima operativni sistem, pri čemu ULT '
  'može dati finiju kontrolu, ali može biti manje efikasan od KLT.')
P('Distraktori su pokazali uspešno ciljanje složenih zabluda. Opcija da i ULT i KLT zahtevaju '
  'eksplicitne mehanizme sinhronizacije deluje vrlo uverljivo, jer je sinhronizacija stvarno '
  'važna tema kod niti, ali pogrešno uopštava nešto što značajno zavisi od implementacije. Opcija '
  'da jezgro upravlja i jednima i drugima sa jednakom efikasnošću i kontrolom cilja zabludu da '
  'upravljanje na nivou jezgra automatski izjednačava sve osobine niti. Najslabiji distraktor je '
  'tvrdnja da ULT obezbeđuje više jedinica resursa od KLT, jer je suviše tehnička i dovoljno '
  'nejasna da je studenti mogu odbaciti zbog formulacije, a ne zbog stvarnog razumevanja.')
H('Relacioni nivo', 3)
P('Relaciona pitanja traže dublje razumevanje, povezivanje i poređenje pojmova. Kao primer '
  'uzimamo pitanje iz kursa Testiranje softvera: Kako automatizacija testova doprinosi '
  'efikasnosti regresionog testiranja? Tačan odgovor je da automatizacija omogućava brže '
  'izvršavanje i otkrivanje promena u funkcionalnosti, čime se povećava temeljnost regresionih '
  'testova.')
P('Pitanje traži da student poveže dva pojma, šta regresiono testiranje jeste i kako automatizacija '
  'rešava njegove izazove, pa se ispravno svrstava u relacioni nivo. Među distraktorima, opcija '
  'da automatizacija u potpunosti zamenjuje ručno testiranje hvata krajnju zabludu o zameni '
  'ljudske procene, dok opcija da automatizacija poboljšava samo brzinu funkcionalnog testiranja '
  'pravi veštačku granicu između funkcionalnog i regresionog testiranja. Opcija da se '
  'automatizacija bavi pravljenjem novih test slučajeva je najslabija, jer se ne poklapa ni sa '
  'jednom čestom zabludom, pošto većina studenata zna da automatizacija izvršava testove, a ne da '
  'ih pravi, pa je lako odbacuju.')
H('Prošireno apstraktni nivo', 3)
P('Prošireno apstraktna pitanja traže uopštavanje i primenu naučenog u novom, složenijem '
  'kontekstu. Kao primer uzimamo pitanje: U distribuiranom računarskom okruženju, ako glavni '
  'proces stvara više procesa dece na različitim čvorovima, koji mehanizam obezbeđuje da svi '
  'procesi deca dobiju ispravne sistemske resurse i da se međusobno ne ometaju? Tačan odgovor je '
  'primena protokola međuprocesne komunikacije (IPC) za upravljanje dodelom resursa i '
  'sinhronizacijom.')
P('Pitanje uspešno uzima pojam upravljanja procesima, koji se obično uči u kontekstu jedne '
  'mašine, i traži da se prenese na distribuirano okruženje. Opcija da se svi procesi deca dodele '
  'jednom čvoru radi centralizovanog upravljanja dobro cilja zabludu da je centralizacija rešenje, '
  'iako ona poništava sam smisao distribucije. Opcija o sinhronizaciji na nivou niti unutar svakog '
  'čvora prepoznaje delimičnu istinu, jer rešava samo probleme unutar čvora, a ne širi '
  'distribuirani problem. Opcija o povećanju prioriteta roditeljskog procesa nad decom je suviše '
  'očigledno netačna, jer povećanje prioriteta logički nema veze sa dodelom resursa ili '
  'sprečavanjem ometanja između čvorova, pa se lako eliminiše.')
P('Pregled po nivoima otkriva jasan obrazac koji je uočen i u početnom radu: kako složenost po '
  'SOLO taksonomiji raste, kvalitet pitanja opada. Niža pitanja, unistrukturalna i '
  'multistrukturalna, po pravilu se dobro poklapaju sa zadatim nivoom i imaju uverljive '
  'distraktore. Viša pitanja, relaciona i naročito prošireno apstraktna, pokazuju veće probleme: '
  'tačan odgovor ponekad sadrži tehničke netačnosti ili pojednostavljenja, a među distraktorima '
  'se javljaju i oni koji su trivijalno netačni ili ne odražavaju stvarne greške u rezonovanju. '
  'Kao ponavljajući problem javlja se i sličnost distraktora, gde dve opcije prenose istu '
  'pogrešnu ideju, čime se smanjuje moć pitanja da razlikuje studente. Upravo zbog ovog obrasca '
  'izgrađen je merni i korektivni sloj koji je predmet ostatka ovog poglavlja, a kvantitativne '
  'mere koje slede pokazuju koliko se taj problem može ublažiti jačim modelom i ciljanim '
  'proverama.')

H('4.3. Prelazak sa lokalnog na globalni model', 2)
P('Najvažniji rezultat je poređenje lokalnog i globalnog modela na istom skupu pitanja i istom '
  'sloju provera. Sledeća slika prikazuje kontrolne table kvaliteta za sve tri lekcije, gde leva '
  'kolona odgovara lokalnom modelu, a desna globalnom. Razlika je izrazita kod svih provera '
  'koje traže zaključivanje.')
add_figure(IMG,
           'Kontrolne table kvaliteta za tri lekcije. Leva kolona je lokalni model '
           '(Qwen 2.5 14B), desna je globalni model (Claude Haiku 4.5).', 6.2)
P('Tabela 3 sažima ključne mere za sve tri lekcije, pre i posle prelaska na globalni model. '
  'Vrednosti su zaokružene onako kako ih sistem prikazuje.')
table(
    ['Mera', 'Procesi L', 'Procesi G', 'Niti L', 'Niti G', 'Konkur. L', 'Konkur. G'],
    [
        ['Pokrivenost pojmova', '58,3%', '54,6%', '58,0%', '63,9%', '41,1%', '50,6%'],
        ['Haladyna prosek', '95,3', '95,4', '96,5', '95,4', '96,0', '95,9'],
        ['SOLO kapa', '0,63', '0,78', '0,64', '0,78', '0,27', '0,68'],
        ['CoVe potvrđeno', '10%', '38,9%', '14,3%', '59,2%', '18,8%', '41,1%'],
        ['Rešivost (mean p)', '0,66', '0,88', '0,71', '0,94', '0,78', '0,91'],
        ['Rešivost iz teksta', '0%', '50%', '0%', '69%', '0%', '64,2%'],
        ['IOC indeks', '-0,01', '0,69', '-0,04', '0,66', '0,02', '0,62'],
        ['Čitljivost (FK)', '17,4', '19,1', '13,5', '16,4', '16,1', '21,3'],
        ['Dvosmislenost', '52,5%', '4,4%', '48,2%', '7%', '51,6%', '3,2%'],
        ['Gram. ujednačenost', '68,8%', '92,2%', '67,9%', '87,3%', '64,1%', '90,5%'],
        ['Uverljivost distr.', '3,45', '3,69', '3,37', '3,58', '3,47', '3,51'],
    ],
)
caption('Tabela 3. Mere kvaliteta po lekcijama, lokalni (L) i globalni (G) model.')
P('Iz tabele se vidi nekoliko jasnih obrazaca. Slaganje sa SOLO nivoom, mereno Koenovim '
  'koeficijentom, sa lokalnim modelom je bilo od 0,27 do 0,64, što je u rasponu od slabog do '
  'umerenog, dok je sa globalnim modelom poraslo na 0,68 do 0,78, što se tumači kao značajno '
  'slaganje. Lanac provere, koji meri koliko je tačnih odgovora zaista potkrepljeno materijalom, '
  'sa lokalnim modelom je potvrđivao svega 10 do 18,8 procenata pitanja, a sa globalnim 38,9 do '
  '59,2 procenata. Provera rešivosti iz teksta je sa lokalnim modelom bila na nuli za sve tri '
  'lekcije, jer model nije umeo da odgovori na pitanje bez gledanja u opcije, dok je sa '
  'globalnim modelom prošla 50 do 69 procenata pitanja. Indeks podudaranja sa ciljem je sa '
  'lokalnim modelom bio oko nule ili negativan, što znači da pitanja nisu merila ono što je '
  'trebalo, a sa globalnim je porastao na vrednosti oko 0,6 do 0,7, koje se smatraju '
  'prihvatljivim. Dvosmislenost je sa lokalnog modela, gde je oko polovine pitanja bilo '
  'označeno kao dvosmisleno, pala na svega nekoliko procenata. Gramatička ujednačenost opcija '
  'je porasla sa oko dve trećine na preko 87 procenata.')
P('Nasuprot tome, deterministične mere koje ne zavise mnogo od modela ostale su slične. '
  'Prosečna ocena po Haladyna pravilima je u oba slučaja blizu 95 od 100, jer ta pravila '
  'proverava programski mehanizam, a ne model. Čitljivost i pokrivenost pojmova se razlikuju, '
  'ali bez jasnog smera, pošto zavise i od toga kako je izgrađena struktura kursa. Uverljivost '
  'distraktora je blago porasla. Iz svega ovoga sledi jasan zaključak: lokalni model od '
  'četrnaest milijardi parametara nije bio dovoljan za sloj koji sudi o pitanjima, pa je za '
  'pedagoški smislenu proveru bio neophodan jači model. Same provere, kao programski okvir, '
  'radile su isto u oba slučaja, ali je kvalitet suda zavisio od modela.')

P('Naredne tri slike vizuelno sažimaju ovu razliku. Prva daje grupisani prikaz ključnih mera u '
  'proseku za tri lekcije, druga isti odnos prikazuje kao radarski profil, a treća kao toplotnu '
  'mapu poboljšanja po lekciji i meri.')
FIG('fig1_local_global.png',
    'Ključne mere sloja provera, lokalni naspram globalnog modela (prosek tri lekcije). Kod svih '
    'mera viša vrednost je bolja; dvosmislenost je prikazana kao jasnoća, odnosno 100 minus stopa '
    'dvosmislenosti.', 6.2)
FIG('fig3_radar.png',
    'Radarski profil kvaliteta. Površina globalnog modela obuhvata površinu lokalnog na svim '
    'osama.', 4.8)
FIG('fig9_improvement_heatmap.png',
    'Toplotna mapa poboljšanja: razlika između globalnog i lokalnog modela, po meri i lekciji. '
    'Tamnije polje znači veći porast.', 5.8)

H('4.4. Poboljšanje lanca provere', 2)
P('Tokom rada uočeno je da lanac provere često označava pitanja kao neodređena ne zato što su '
  'loša, nego zato što mu je dat premali kontekst. U prvobitnoj verziji se tačnost proveravala '
  'samo u odnosu na jedan citat iz materijala, što je za mnoga pitanja premalo. Provera je '
  'izmenjena tako da koristi širi kontekst, odnosno tekst cele sekcije ili nastavnog objekta uz '
  'citat. Posle te izmene, na lekciji Procesi je udeo potvrđenih pitanja porastao sa 38,9 na '
  '54,4 procenta, dakle za petnaest procentnih poena, bez ikakve promene u samim pitanjima. To '
  'pokazuje da se deo ranije označenih pitanja nije zaista loš, nego je proveri jednostavno '
  'nedostajao kontekst.')
FIG('fig2_cove.png',
    'Udeo potvrđenih odgovora po lekciji. Za lekciju Procesi prikazana je i vrednost posle '
    'davanja šireg konteksta, koja raste sa 38,9 na 54,4 procenta.', 5.8)

H('4.5. Spoljna provera na referentnom skupu EduQG', 2)
P('Da bi se proverilo da li su sami merni instrumenti dobro podešeni, sloj provera je pušten i '
  'nad pitanjima iz skupa EduQG, koja su pisali stručnjaci. Ideja je jednostavna: ako su provere '
  'dobro kalibrisane, one bi trebalo da propuste veliku većinu stručnjačkih pitanja, jer su to '
  'po pretpostavci dobra pitanja. Korišćen je stratifikovan uzorak od 149 pitanja iz dvanaest '
  'udžbenika, a kao model za sud je korišćen Claude Haiku 4.5, isti kao u produkciji. Tabela 4 '
  'prikazuje koliko često je svaka provera označila stručnjačko pitanje kao problematično. Niža '
  'vrednost je bolja, jer znači manje lažnih uzbuna.')
table(
    ['Provera', 'Udeo označenih (niže je bolje)'],
    [
        ['Haladyna pravila (greška)', '0,0%'],
        ['Rešivost (p < 0,5)', '4,7% (prosečno p = 0,94)'],
        ['Gramatička ujednačenost (tačan je odudara)', '4,7%'],
        ['Uverljivost distraktora (ispod 2,5)', '0,7% (prosek 3,92 od 5)'],
        ['Dvosmislenost', '24,8%'],
        ['Lanac provere (nije potvrđeno)', '42,3%'],
    ],
)
caption('Tabela 4. Udeo stručnjačkih pitanja koje je svaka provera označila (specifičnost).')
P('Najbolje se pokazala provera rešivosti, koja je propustila gotovo sva stručnjačka pitanja, a '
  'uz to, na namerno pokvarenim pitanjima, kod kojih je tačan odgovor zamenjen netačnim, uspela '
  'da prepozna 98 procenata grešaka. To je znak dobro kalibrisane provere, jer i propušta dobra '
  'pitanja i hvata loša. Provere zasnovane na pravilima, gramatici i uverljivosti takođe su '
  'dobro podešene. Sa druge strane, lanac provere je previše strog, jer je označio 42,3 procenta '
  'stručnjačkih pitanja, mada je gotovo sve to bilo neodređeno, a ne stvarno oboreno. Davanjem '
  'šireg konteksta taj udeo pada na 35,6 procenata, što potvrđuje da je deo problema upravo u '
  'nedostatku konteksta, ali pokazuje i da je provera po prirodi stroga za pitanja koja se '
  'oslanjaju na predznanje. Provera dvosmislenosti je pokazala umerenu strogost, sa 24,8 '
  'procenata, pri čemu je pregled primera pokazao da se delom radi o stvarno kontekstualno '
  'zavisnim pitanjima, a delom o lažnim uzbunama.')
FIG('fig5_eduqg_calibration.png',
    'Kalibracija provera na skupu EduQG. Levo je udeo stručnjačkih pitanja koje svaka provera '
    'označi (niže je bolje), desno je udeo uhvaćenih namerno pokvarenih pitanja (više je bolje).',
    6.4)
P('Pored toga, na istom uzorku je upoređeno generisanje distraktora sa stručnjačkim. Sistem je '
  'za svako pitanje napravio tri distraktora, koji su zatim upoređeni sa tri stručnjačka. '
  'Leksičko poklapanje sa stručnjačkim distraktorima je očekivano nisko, oko 12 procenata za '
  'tačno poklapanje i oko 19,5 procenata za delimično, jer postoji mnogo validnih distraktora '
  'koji se međusobno razlikuju. Po meri uverljivosti, distraktori sistema su ocenjeni neznatno '
  'više od stručnjačkih, 4,07 prema 3,95, ali ta brojka traži oprez. Pregled je otkrio da '
  'sistem u oko 11 procenata slučajeva napravi distraktor koji je zapravo parafraza ili '
  'preuređenje tačnog odgovora, a da ga provera uverljivosti pri tom ne kažnjava. Taj nalaz je '
  'iskorišćen za dve konkretne ispravke: u prompt za distraktore dodato je izričito pravilo da '
  'distraktor ne sme biti parafraza tačnog odgovora, a u mehanizam za pravila dodata je '
  'deterministička provera koja označava distraktor čiji je skup reči istovetan tačnom odgovoru. '
  'Time je spoljni referentni skup poslužio i kao alat za otkrivanje skrivene mane i kao osnova '
  'za njeno otklanjanje.')

FIG('fig6_eduqg_distractors.png',
    'Poređenje distraktora sa stručnjačkim. Levo je prosečna uverljivost, desno udeo obnovljenih '
    'stručnjačkih distraktora.', 5.6)

H('4.6. Osvrt na rezultate i ograničenja', 2)
P('Rezultati daju odgovore na sva četiri istraživačka pitanja. Na prvo pitanje, o usklađenosti '
  'sa SOLO nivoom, odgovor je da je usklađenost sa globalnim modelom značajna, sa Koenovim '
  'koeficijentom od 0,68 do 0,78. Na drugo pitanje, o kvalitetu distraktora, odgovor je da su '
  'distraktori uverljivi po meri ocene, ali da je postojala mana parafraziranja tačnog odgovora, '
  'koja je sada otklonjena. Na treće pitanje, o pokrivenosti, odgovor je da pitanja pokrivaju '
  'oko polovine do dve trećine pojmova po lekciji, što ostavlja prostora za više pitanja po '
  'pojmu. Na četvrto pitanje, o tačnosti i ispravnosti, odgovor je da je provera rešivosti vrlo '
  'pouzdana, dok je lanac provere strog, ali ga širi kontekst poboljšava.')
P('Glavno ograničenje je veličina studije slučaja, koja obuhvata jedan kurs i tri lekcije, pa '
  'brojke daju jasan smer, ali ne i uske intervale poverenja. Drugo ograničenje je to što su '
  'rezultati za lokalni i globalni model snimljeni u različitim trenucima, pa pojedine '
  'deterministične mere zavise i od međuvremenskih izmena u izgradnji strukture kursa. Treće, '
  'spoljni skup EduQG je na engleskom jeziku i iz drugih oblasti, dok je studija slučaja na '
  'srpskom i iz oblasti operativnih sistema, pa brojke o generisanju treba čitati u tom svetlu. '
  'Uprkos tome, glavni nalaz je jasan i dosledan: sistem sa slojem za kontrolu kvaliteta i '
  'jačim modelom daje pitanja koja su merljivo bolja od polazne verzije.')

# ============================ 5. PROTOTIP ============================
H('5. Prototip softverske aplikacije', 1)
P('Sistem je dostupan kroz veb aplikaciju koja vodi nastavnika kroz ceo postupak, od unosa '
  'materijala do gotove provere. Korisnički deo je napravljen u tehnologiji React, a serverski '
  'u okviru Flask. Aplikacija je organizovana u nekoliko celina.')
P('U celini za upravljanje kursevima i lekcijama nastavnik pravi kurs, dodaje lekcije i učitava '
  'materijale u PDF formatu. Nakon učitavanja, sistem raščlanjuje materijal i prikazuje '
  'izvučenu strukturu, odnosno sekcije i nastavne objekte, koje nastavnik može da pregleda i '
  'po potrebi ispravi. U celini za pregled sadržaja prikazuje se i ontologija, a moguće je '
  'postavljati i upite jezikom SPARQL nad izgrađenim grafom znanja.')
P('U celini za generisanje pitanja nastavnik pokreće pravljenje pitanja po SOLO nivoima, a '
  'napravljena pitanja se skupljaju u banci pitanja, gde se mogu pregledati, izmeniti ili '
  'obrisati. Iz banke se zatim u celini za sastavljanje kviza bira koliko pitanja po kom nivou '
  'ulazi u proveru, a u celini za rešavanje kviza provera se može i odraditi.')
P('Najvažniji deo za ovaj rad je kontrolna tabla kvaliteta, koja jednim klikom pokreće svih '
  'dvanaest provera nad pitanjima jedne lekcije i prikazuje rezultate u obliku kartica, sa '
  'ocenom i bojom koja odmah pokazuje da li je mera dobra. Pošto neke provere traže mnogo poziva '
  'modela i mogu dugo da traju, rezultati se keširaju, pa se pri ponovnom otvaranju odmah '
  'prikazuju, bez ponovnog računanja. Dodata je i posebna stranica za poređenje sa referentnim '
  'skupom EduQG, koja prikazuje stručnjačka pitanja i pored svakog pokazuje kako je svaka '
  'provera prošla, čime se na istom mestu vidi kalibracija samih mernih instrumenata. Aplikacija '
  'podržava i prevođenje sadržaja, pa se pitanja mogu pripremiti na više jezika.')

# ============================ 6. ZAKLJUCAK ============================
H('6. Zaključak', 1)
P('U ovom radu je predstavljen sistem za automatsko generisanje pitanja za proveru znanja po '
  'SOLO taksonomiji, koji spaja velike jezičke modele sa ontološkim opisom gradiva i sa slojem '
  'za kontrolu kvaliteta zasnovanim na pedagoškoj i psihometrijskoj literaturi. Za razliku od '
  'početne verzije sistema, kod koje je kvalitet procenjivan kvalitativno, ovde je kvalitet '
  'izmeren brojkama, po lekcijama i zbirno, i to kroz dvanaest provera povezanih sa četiri '
  'istraživačka pitanja.')
P('Glavni nalaz je da sloj za kontrolu kvaliteta u velikoj meri zavisi od sposobnosti modela da '
  'sudi, pa je prelazak sa lokalnog modela od četrnaest milijardi parametara na jači model u '
  'oblaku doneo izrazito poboljšanje na svim merama koje traže zaključivanje, od slaganja sa '
  'SOLO nivoom i indeksa podudaranja sa ciljem, preko lanca provere i rešivosti iz teksta, do '
  'smanjenja dvosmislenosti i bolje gramatičke ujednačenosti opcija. Pored toga, pokazano je da '
  'se ciljanim izmenama kvalitet dalje poboljšava, na primer davanjem šireg konteksta lancu '
  'provere, čime je udeo potvrđenih pitanja na jednoj lekciji porastao sa 38,9 na 54,4 procenta. '
  'Najzad, proverom na spoljnom referentnom skupu EduQG kalibrisani su sami merni instrumenti, '
  'pri čemu se provera rešivosti pokazala kao veoma pouzdana, a otkrivena je i otklonjena '
  'skrivena mana u generisanju distraktora.')
P('Budući rad ima nekoliko pravaca. Prvi je veća studija evaluacije, sa više kurseva i lekcija, '
  'da bi se dobile uže ocene i intervali poverenja. Drugi je dalje podešavanje pojedinih '
  'provera, pre svega praga strogosti lanca provere i provere dvosmislenosti. Treći je proba '
  'sa drugim modelima i merenje odnosa brzine, cene i kvaliteta, jer je sa sve većim brojem '
  'provera vreme obrade osetno poraslo, pa se postavlja pitanje da li bi brži model u oblaku '
  'dao i bolji kvalitet i kraće vreme. Četvrti je proširenje sistema izvan pitanja višestrukog '
  'izbora, ka pitanjima sa kratkim odgovorom i esejskim pitanjima.')

# ============================ ZAHVALNICA ============================
H('Zahvalnica', 1)
P('Ovo istraživanje je podržalo Ministarstvo nauke, tehnološkog razvoja i inovacija (ugovor br. '
  '451-03-34/2026-03/200156) i Fakultet tehničkih nauka Univerziteta u Novom Sadu kroz projekat '
  'naučnog i umetničkog istraživačkog rada istraživača u nastavnim i saradničkim zvanjima na '
  'Fakultetu tehničkih nauka.')

# ============================ LITERATURA ============================
H('Literatura', 1)
refs = [
 'Chan, C.C., Tsui, M.S., Chan, M.Y. and Hong, J.H., 2002. Applying the structure of the '
 'observed learning outcomes (SOLO) taxonomy on student\'s learning outcomes: An empirical '
 'study. Assessment & Evaluation in Higher Education, 27(6), pp. 511-527.',
 'Smith, B., 2012. Ontology. The furniture of the world, pp. 47-68.',
 'Mitkov, R., 2003. Computer-aided generation of multiple-choice tests. In Proceedings of the '
 'HLT-NAACL 03 workshop on Building educational applications using natural language processing '
 '(pp. 17-22).',
 'Kasneci, E. i dr., 2023. ChatGPT for good? On opportunities and challenges of large language '
 'models for education. Learning and individual differences, 103, p. 102274.',
 'Lister, R., Simon, B., Thompson, E., Whalley, J.L. and Prasad, C., 2006. Not seeing the '
 'forest for the trees: novice programmers and the SOLO taxonomy. ACM SIGCSE Bulletin, 38(3), '
 'pp. 118-122.',
 'Liang, C., Yang, X., Dave, N., Wham, D., Pursel, B. and Giles, C.L., 2018. Distractor '
 'generation for multiple choice questions using learning to rank. In Proceedings of the '
 'thirteenth workshop on innovative use of NLP for building educational applications (pp. '
 '284-290).',
 'Scaria, N., Dharani Chenna, S. and Subramani, D., 2024. Automated educational question '
 'generation at different Bloom\'s skill levels using large language models: Strategies and '
 'evaluation. In International conference on artificial intelligence in education (pp. 165-179). '
 'Springer.',
 'Wei, J., Wang, X., Schuurmans, D., Bosma, M., Xia, F., Chi, E., Le, Q.V. and Zhou, D., 2022. '
 'Chain-of-thought prompting elicits reasoning in large language models. Advances in neural '
 'information processing systems, 35, pp. 24824-24837.',
 'Sweller, J. and Cooper, G.A., 1985. The use of worked examples as a substitute for problem '
 'solving in learning algebra. Cognition and instruction, 2(1), pp. 59-89.',
 'Lewis, P. i dr., 2020. Retrieval-augmented generation for knowledge-intensive NLP tasks. '
 'Advances in neural information processing systems, 33, pp. 9459-9474.',
 'Bitew, S.K., Deleu, J., Develder, C. and Demeester, T., 2023. Distractor generation for '
 'multiple-choice questions with predictive prompting and large language models. In Joint '
 'European Conference on Machine Learning and Knowledge Discovery in Databases (pp. 48-63). '
 'Springer.',
 'Zhang, T., Kishore, V., Wu, F., Weinberger, K.Q. and Artzi, Y., 2019. BERTScore: Evaluating '
 'text generation with BERT. arXiv preprint arXiv:1904.09675.',
 'Chen, C.H. and Shiu, M.F., 2025. KAQG: A knowledge-graph-enhanced RAG for difficulty-'
 'controlled question generation. IEEE Access, 13, pp. 197234-197244.',
 'Haladyna, T.M., Downing, S.M. and Rodriguez, M.C., 2002. A review of multiple-choice '
 'item-writing guidelines for classroom assessment. Applied measurement in education, 15(3), '
 'pp. 309-333.',
 'Dhuliawala, S. i dr., 2024. Chain-of-verification reduces hallucination in large language '
 'models. In Findings of the Association for Computational Linguistics: ACL 2024 (pp. '
 '3563-3578).',
 'Crocker, L. and Algina, J., 1986. Introduction to classical and modern test theory. Holt, '
 'Rinehart and Winston.',
 'Landis, J.R. and Koch, G.G., 1977. The measurement of observer agreement for categorical '
 'data. Biometrics, 33(1), pp. 159-174.',
 'Cohen, J., 1960. A coefficient of agreement for nominal scales. Educational and psychological '
 'measurement, 20(1), pp. 37-46.',
 'Kurdi, G., Leo, J., Parsia, B., Sattler, U. and Al-Emari, S., 2020. A systematic review of '
 'automatic question generation for educational purposes. International journal of artificial '
 'intelligence in education, 30(1), pp. 121-204.',
 'Flesch, R., 1948. A new readability yardstick. Journal of applied psychology, 32(3), p. 221.',
 'Kincaid, J.P., Fishburne Jr, R.P., Rogers, R.L. and Chissom, B.S., 1975. Derivation of new '
 'readability formulas for navy enlisted personnel. Research Branch Report 8-75.',
 'Downing, S.M., 2005. The effects of violating standard item writing principles on tests and '
 'students. Advances in health sciences education, 10(2), pp. 133-143.',
 'Sadler, P.M., 1998. Psychometric models of student conceptions in science. Journal of '
 'Research in Science Teaching, 35(3), pp. 265-296.',
 'Hadifar, A., Bitew, S.K., Deleu, J., Develder, C. and Demeester, T., 2023. EduQG: A '
 'multi-format multiple-choice dataset for the educational domain. IEEE Access, 11, pp. '
 '20885-20896.',
]
for i, r in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.add_run(f'[{i}] ').bold = True
    p.add_run(r)

out = os.path.join(os.path.dirname(__file__), 'MasterRad.docx')
doc.save(out)
print('SAVED', out)
print('paragraphs:', len(doc.paragraphs))
